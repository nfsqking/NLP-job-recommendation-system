# -*- coding: utf-8 -*-
"""简历路由模块

处理简历录入、提交、查看、编辑、切换、删除等请求。
支持用户管理最多3份独立简历。
"""

import json
import io
import re
from flask import Blueprint, render_template, redirect, url_for, flash, request, session, jsonify, current_app
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from app import db
from app.models import Resume

try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    from zai import ZhipuAiClient
    GLM_SUPPORT = True
except ImportError:
    GLM_SUPPORT = False

resume = Blueprint('resume', __name__)

MAX_RESUME_COUNT = 3
MAX_FILE_SIZE = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {'pdf', 'docx'}


def parse_list_data(data_list):
    """将列表数据转换为JSON字符串存储"""
    if not data_list:
        return ''
    if isinstance(data_list, list):
        return json.dumps(data_list, ensure_ascii=False)
    return data_list


def get_list_data(json_str):
    """将JSON字符串转换为列表数据"""
    if not json_str:
        return []
    try:
        return json.loads(json_str)
    except:
        return []


def prepare_resume_data(resume):
    """准备简历数据用于前端回显"""
    if not resume:
        return {}
    
    education_list = get_list_data(resume.education)
    internship_list = get_list_data(resume.internship)
    work_list = get_list_data(resume.work)
    project_list = get_list_data(resume.project)
    
    return {
        'education': education_list if education_list else [{'school': '', 'major': ''}],
        'internship': internship_list,
        'work': work_list,
        'project': project_list
    }


def get_current_resume_id():
    """获取当前选中的简历ID，如果没有则返回活跃简历ID"""
    resume_id = session.get('current_resume_id')
    if resume_id:
        resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first()
        if resume:
            return resume_id
    
    active_resume = Resume.query.filter_by(user_id=current_user.id, is_active=True).first()
    if active_resume:
        session['current_resume_id'] = active_resume.id
        return active_resume.id
    
    first_resume = Resume.query.filter_by(user_id=current_user.id).first()
    if first_resume:
        if Resume.query.filter_by(user_id=current_user.id, is_active=True).count() == 0:
            first_resume.is_active = True
            db.session.commit()
        session['current_resume_id'] = first_resume.id
        return first_resume.id
    
    return None


def get_current_resume():
    """获取当前选中的简历对象"""
    resume_id = get_current_resume_id()
    if resume_id:
        return Resume.query.get(resume_id)
    return None


def set_current_resume(resume_id):
    """设置当前选中的简历"""
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first()
    if not resume:
        return False
    
    Resume.query.filter_by(user_id=current_user.id).update({'is_active': False})
    resume.is_active = True
    db.session.commit()
    
    session['current_resume_id'] = resume_id
    return True


@resume.route('/list')
@login_required
def list_resumes():
    """简历列表页面"""
    resumes = Resume.query.filter_by(user_id=current_user.id).order_by(Resume.created_at.desc()).all()
    current_resume_id = get_current_resume_id()
    
    return render_template('resume/resume_list.html', 
                          resumes=resumes, 
                          current_resume_id=current_resume_id,
                          max_count=MAX_RESUME_COUNT)


@resume.route('/input', methods=['GET', 'POST'])
@login_required
def input_resume():
    """简历录入页面 - 创建新简历"""
    resume_count = Resume.query.filter_by(user_id=current_user.id).count()
    if resume_count >= MAX_RESUME_COUNT:
        flash(f'最多只能创建 {MAX_RESUME_COUNT} 份简历', 'warning')
        return redirect(url_for('resume.list_resumes'))
    
    if request.method == 'POST':
        resume_name = request.form.get('resume_name', '').strip()
        if not resume_name:
            resume_name = f'简历 {resume_count + 1}'
        
        education_schools = request.form.getlist('education_school')
        education_majors = request.form.getlist('education_major')
        
        internship_companies = request.form.getlist('internship_company')
        internship_positions = request.form.getlist('internship_position')
        internship_contents = request.form.getlist('internship_content')
        
        work_companies = request.form.getlist('work_company')
        work_positions = request.form.getlist('work_position')
        work_contents = request.form.getlist('work_content')
        
        project_names = request.form.getlist('project_name')
        project_contents = request.form.getlist('project_content')
        
        skills = request.form.get('skills', '').strip()
        self_evaluation = request.form.get('self_evaluation', '').strip()
        
        has_education_error = False
        for school, major in zip(education_schools, education_majors):
            if not school.strip() or not major.strip():
                has_education_error = True
                break
        
        if has_education_error:
            flash('请填写完整的教育经历（学校名称和专业为必填项）', 'danger')
            resume_data = {
                'education': [{'school': s, 'major': m} for s, m in zip(education_schools, education_majors)],
                'internship': [{'company': c, 'position': p, 'content': ct} for c, p, ct in zip(internship_companies, internship_positions, internship_contents)],
                'work': [{'company': c, 'position': p, 'content': ct} for c, p, ct in zip(work_companies, work_positions, work_contents)],
                'project': [{'name': n, 'content': c} for n, c in zip(project_names, project_contents)]
            }
            return render_template('resume/resume_input.html', resume=None, resume_data=resume_data, resume_name=resume_name)
        
        if not skills:
            flash('个人技能不能为空', 'danger')
            resume_data = {
                'education': [{'school': s, 'major': m} for s, m in zip(education_schools, education_majors)],
                'internship': [{'company': c, 'position': p, 'content': ct} for c, p, ct in zip(internship_companies, internship_positions, internship_contents)],
                'work': [{'company': c, 'position': p, 'content': ct} for c, p, ct in zip(work_companies, work_positions, work_contents)],
                'project': [{'name': n, 'content': c} for n, c in zip(project_names, project_contents)]
            }
            return render_template('resume/resume_input.html', resume=None, resume_data=resume_data, resume_name=resume_name)
        
        if not self_evaluation:
            flash('个人评价不能为空', 'danger')
            resume_data = {
                'education': [{'school': s, 'major': m} for s, m in zip(education_schools, education_majors)],
                'internship': [{'company': c, 'position': p, 'content': ct} for c, p, ct in zip(internship_companies, internship_positions, internship_contents)],
                'work': [{'company': c, 'position': p, 'content': ct} for c, p, ct in zip(work_companies, work_positions, work_contents)],
                'project': [{'name': n, 'content': c} for n, c in zip(project_names, project_contents)]
            }
            return render_template('resume/resume_input.html', resume=None, resume_data=resume_data, resume_name=resume_name)
        
        education_list = [{'school': s.strip(), 'major': m.strip()} for s, m in zip(education_schools, education_majors) if s.strip() or m.strip()]
        internship_list = [{'company': c.strip(), 'position': p.strip(), 'content': ct.strip()} for c, p, ct in zip(internship_companies, internship_positions, internship_contents) if c.strip() or p.strip() or ct.strip()]
        work_list = [{'company': c.strip(), 'position': p.strip(), 'content': ct.strip()} for c, p, ct in zip(work_companies, work_positions, work_contents) if c.strip() or p.strip() or ct.strip()]
        project_list = [{'name': n.strip(), 'content': c.strip()} for n, c in zip(project_names, project_contents) if n.strip() or c.strip()]
        
        is_first_resume = Resume.query.filter_by(user_id=current_user.id).count() == 0
        
        new_resume = Resume(
            user_id=current_user.id,
            resume_name=resume_name,
            is_active=is_first_resume,
            education=parse_list_data(education_list),
            internship=parse_list_data(internship_list),
            work=parse_list_data(work_list),
            project=parse_list_data(project_list),
            skills=skills,
            self_evaluation=self_evaluation
        )
        db.session.add(new_resume)
        db.session.commit()
        
        if is_first_resume:
            session['current_resume_id'] = new_resume.id
        
        flash('简历创建成功', 'success')
        return redirect(url_for('resume.list_resumes'))
    
    resume_data = prepare_resume_data(None)
    return render_template('resume/resume_input.html', resume=None, resume_data=resume_data, resume_name='')


@resume.route('/edit/<int:resume_id>', methods=['GET', 'POST'])
@login_required
def edit_resume(resume_id):
    """简历编辑页面"""
    existing_resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first()
    
    if not existing_resume:
        flash('简历不存在', 'danger')
        return redirect(url_for('resume.list_resumes'))
    
    if request.method == 'POST':
        resume_name = request.form.get('resume_name', '').strip()
        if not resume_name:
            resume_name = existing_resume.resume_name
        
        education_schools = request.form.getlist('education_school')
        education_majors = request.form.getlist('education_major')
        
        internship_companies = request.form.getlist('internship_company')
        internship_positions = request.form.getlist('internship_position')
        internship_contents = request.form.getlist('internship_content')
        
        work_companies = request.form.getlist('work_company')
        work_positions = request.form.getlist('work_position')
        work_contents = request.form.getlist('work_content')
        
        project_names = request.form.getlist('project_name')
        project_contents = request.form.getlist('project_content')
        
        skills = request.form.get('skills', '').strip()
        self_evaluation = request.form.get('self_evaluation', '').strip()
        
        has_education_error = False
        for school, major in zip(education_schools, education_majors):
            if not school.strip() or not major.strip():
                has_education_error = True
                break
        
        if has_education_error:
            flash('请填写完整的教育经历（学校名称和专业为必填项）', 'danger')
            resume_data = {
                'education': [{'school': s, 'major': m} for s, m in zip(education_schools, education_majors)],
                'internship': [{'company': c, 'position': p, 'content': ct} for c, p, ct in zip(internship_companies, internship_positions, internship_contents)],
                'work': [{'company': c, 'position': p, 'content': ct} for c, p, ct in zip(work_companies, work_positions, work_contents)],
                'project': [{'name': n, 'content': c} for n, c in zip(project_names, project_contents)]
            }
            return render_template('resume/resume_input.html', resume=existing_resume, resume_data=resume_data, resume_name=resume_name)
        
        if not skills:
            flash('个人技能不能为空', 'danger')
            resume_data = {
                'education': [{'school': s, 'major': m} for s, m in zip(education_schools, education_majors)],
                'internship': [{'company': c, 'position': p, 'content': ct} for c, p, ct in zip(internship_companies, internship_positions, internship_contents)],
                'work': [{'company': c, 'position': p, 'content': ct} for c, p, ct in zip(work_companies, work_positions, work_contents)],
                'project': [{'name': n, 'content': c} for n, c in zip(project_names, project_contents)]
            }
            return render_template('resume/resume_input.html', resume=existing_resume, resume_data=resume_data, resume_name=resume_name)
        
        if not self_evaluation:
            flash('个人评价不能为空', 'danger')
            resume_data = {
                'education': [{'school': s, 'major': m} for s, m in zip(education_schools, education_majors)],
                'internship': [{'company': c, 'position': p, 'content': ct} for c, p, ct in zip(internship_companies, internship_positions, internship_contents)],
                'work': [{'company': c, 'position': p, 'content': ct} for c, p, ct in zip(work_companies, work_positions, work_contents)],
                'project': [{'name': n, 'content': c} for n, c in zip(project_names, project_contents)]
            }
            return render_template('resume/resume_input.html', resume=existing_resume, resume_data=resume_data, resume_name=resume_name)
        
        education_list = [{'school': s.strip(), 'major': m.strip()} for s, m in zip(education_schools, education_majors) if s.strip() or m.strip()]
        internship_list = [{'company': c.strip(), 'position': p.strip(), 'content': ct.strip()} for c, p, ct in zip(internship_companies, internship_positions, internship_contents) if c.strip() or p.strip() or ct.strip()]
        work_list = [{'company': c.strip(), 'position': p.strip(), 'content': ct.strip()} for c, p, ct in zip(work_companies, work_positions, work_contents) if c.strip() or p.strip() or ct.strip()]
        project_list = [{'name': n.strip(), 'content': c.strip()} for n, c in zip(project_names, project_contents) if n.strip() or c.strip()]
        
        existing_resume.resume_name = resume_name
        existing_resume.education = parse_list_data(education_list)
        existing_resume.internship = parse_list_data(internship_list)
        existing_resume.work = parse_list_data(work_list)
        existing_resume.project = parse_list_data(project_list)
        existing_resume.skills = skills
        existing_resume.self_evaluation = self_evaluation
        
        db.session.commit()
        flash('简历更新成功', 'success')
        return redirect(url_for('resume.list_resumes'))
    
    resume_data = prepare_resume_data(existing_resume)
    return render_template('resume/resume_input.html', resume=existing_resume, resume_data=resume_data, resume_name=existing_resume.resume_name)


@resume.route('/detail/<int:resume_id>')
@login_required
def detail_resume(resume_id):
    """简历查看页面"""
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first()
    
    if not resume:
        flash('简历不存在', 'danger')
        return redirect(url_for('resume.list_resumes'))
    
    education_list = get_list_data(resume.education)
    internship_list = get_list_data(resume.internship)
    work_list = get_list_data(resume.work)
    project_list = get_list_data(resume.project)
    
    current_resume_id = get_current_resume_id()
    
    return render_template('resume/resume_detail.html', 
                          resume=resume, 
                          education_list=education_list,
                          internship_list=internship_list,
                          work_list=work_list,
                          project_list=project_list,
                          current_resume_id=current_resume_id)


@resume.route('/switch/<int:resume_id>')
@login_required
def switch_resume(resume_id):
    """切换当前使用的简历"""
    if set_current_resume(resume_id):
        flash('已切换到选中的简历', 'success')
    else:
        flash('切换失败，简历不存在', 'danger')
    return redirect(url_for('resume.list_resumes'))


@resume.route('/delete/<int:resume_id>')
@login_required
def delete_resume(resume_id):
    """删除简历"""
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first()
    
    if not resume:
        flash('简历不存在', 'danger')
        return redirect(url_for('resume.list_resumes'))
    
    was_active = resume.is_active
    db.session.delete(resume)
    db.session.commit()
    
    if was_active:
        remaining = Resume.query.filter_by(user_id=current_user.id).first()
        if remaining:
            remaining.is_active = True
            session['current_resume_id'] = remaining.id
            db.session.commit()
        else:
            session.pop('current_resume_id', None)
    
    flash('简历已删除', 'success')
    return redirect(url_for('resume.list_resumes'))


@resume.route('/api/current')
@login_required
def api_get_current_resume():
    """获取当前简历信息API"""
    resume = get_current_resume()
    if not resume:
        return jsonify({
            'success': False,
            'message': '暂无简历'
        })
    
    return jsonify({
        'success': True,
        'resume': resume.to_dict()
    })


@resume.route('/api/list')
@login_required
def api_list_resumes():
    """获取简历列表API"""
    resumes = Resume.query.filter_by(user_id=current_user.id).order_by(Resume.created_at.desc()).all()
    current_resume_id = get_current_resume_id()
    
    return jsonify({
        'success': True,
        'resumes': [r.to_dict() for r in resumes],
        'current_resume_id': current_resume_id
    })


def allowed_file(filename):
    """检查文件扩展名是否合法"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_stream):
    """从PDF文件中提取文本"""
    if not PDF_SUPPORT:
        raise Exception("PDF解析库未安装")
    
    try:
        text_parts = []
        with pdfplumber.open(file_stream) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return '\n'.join(text_parts)
    except Exception as e:
        raise Exception(f"PDF解析失败: {str(e)}")


def extract_text_from_docx(file_stream):
    """从DOCX文件中提取文本"""
    if not DOCX_SUPPORT:
        raise Exception("DOCX解析库未安装")
    
    try:
        doc = Document(file_stream)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' '.join(row_text))
        
        return '\n'.join(text_parts)
    except Exception as e:
        raise Exception(f"DOCX解析失败: {str(e)}")


def extract_resume_info_with_glm(text, api_key):
    """使用GLM-4.7-FLASH提取简历结构化信息"""
    if not GLM_SUPPORT:
        raise Exception("GLM SDK未安装")
    
    prompt = f"""你是专业简历信息提取助手，严格根据用户提供的简历原文，提取信息并输出**严格JSON**，**禁止编造、禁止脑补、禁止空字段填空、禁止返回解释**。
字段如下：
{{
  "resume_name": "简历名称，字符串",
  "education": "教育经历JSON数组，每项包含school和major字段",
  "internship": "实习经历JSON数组，每项包含company、position、content字段",
  "work": "工作经历JSON数组，每项包含company、position、content字段",
  "project": "项目经历JSON数组，每项包含name、content字段",
  "skills": "技能，必填，多技能用分号分隔",
  "self_evaluation": "自我评价，必填，字符串"
}}

规则：
- education、internship、work、project 字段必须是JSON数组格式
- 无信息则返回空数组 []，不要写"无""未知"
- 严格JSON，无多余文字、无Markdown、无解释
- 不要修改原文意思，只提取、整理

示例输出：
{{
  "resume_name": "张三的简历",
  "education": [{{"school": "北京大学", "major": "计算机科学"}}],
  "internship": [{{"company": "腾讯", "position": "后端实习生", "content": "参与XX项目开发"}}],
  "work": [],
  "project": [{{"name": "电商系统", "content": "负责订单模块开发"}}],
  "skills": "Python;Java;MySQL",
  "self_evaluation": "热爱技术，学习能力强"
}}

=== 简历原文 ===
{text}
"""
    
    try:
        client = ZhipuAiClient(api_key=api_key)
        
        response = client.chat.completions.create(
            model="glm-4.7-flash",
            messages=[
                {"role": "user", "content": prompt}
            ],
            thinking={
                "type": "disabled",
            },
            max_tokens=4096,
            temperature=0.1
        )
        
        if response and hasattr(response, 'choices') and len(response.choices) > 0:
            content = response.choices[0].message.content
            return content
        else:
            raise Exception("GLM返回格式异常")
            
    except Exception as e:
        raise Exception(f"GLM调用失败: {str(e)}")


def parse_glm_response(response_text):
    """解析GLM返回的JSON"""
    try:
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            json_str = json_match.group(0)
            data = json.loads(json_str)
            
            result = {
                'resume_name': data.get('resume_name', ''),
                'education': data.get('education', []),
                'internship': data.get('internship', []),
                'work': data.get('work', []),
                'project': data.get('project', []),
                'skills': data.get('skills', ''),
                'self_evaluation': data.get('self_evaluation', '')
            }
            
            return result
        else:
            raise Exception("未找到有效JSON")
    except json.JSONDecodeError as e:
        raise Exception(f"JSON解析失败: {str(e)}")


def convert_to_form_data(extracted_data):
    """将提取的数据转换为表单格式"""
    form_data = {
        'resume_name': extracted_data.get('resume_name', ''),
        'education': extracted_data.get('education', []),
        'internship': extracted_data.get('internship', []),
        'work': extracted_data.get('work', []),
        'project': extracted_data.get('project', []),
        'skills': extracted_data.get('skills', ''),
        'self_evaluation': extracted_data.get('self_evaluation', '')
    }
    
    if not form_data['education']:
        form_data['education'].append({'school': '', 'major': ''})
    
    return form_data


@resume.route('/api/parse', methods=['POST'])
@login_required
def api_parse_resume():
    """解析简历文件API"""
    try:
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'message': '未选择文件'
            })
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({
                'success': False,
                'message': '未选择文件'
            })
        
        if not allowed_file(file.filename):
            return jsonify({
                'success': False,
                'message': '不支持的文件格式，请上传 .pdf 或 .docx 文件'
            })
        
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)
        
        if file_size > MAX_FILE_SIZE:
            return jsonify({
                'success': False,
                'message': f'文件过大，请上传小于 {MAX_FILE_SIZE // (1024*1024)}MB 的文件'
            })
        
        file_ext = file.filename.rsplit('.', 1)[1].lower()
        
        try:
            file_stream = io.BytesIO(file.read())
            
            if file_ext == 'pdf':
                if not PDF_SUPPORT:
                    return jsonify({
                        'success': False,
                        'message': 'PDF解析库未安装，请联系管理员'
                    })
                text = extract_text_from_pdf(file_stream)
            elif file_ext == 'docx':
                if not DOCX_SUPPORT:
                    return jsonify({
                        'success': False,
                        'message': 'DOCX解析库未安装，请联系管理员'
                    })
                text = extract_text_from_docx(file_stream)
            else:
                return jsonify({
                    'success': False,
                    'message': '不支持的文件格式'
                })
            
            if not text or not text.strip():
                return jsonify({
                    'success': False,
                    'message': '未提取到有效文本，请检查文件内容'
                })
            
            print(f"[简历解析] 成功提取文本，长度: {len(text)} 字符")
            
        except Exception as e:
            print(f"[简历解析] 文件解析异常: {str(e)}")
            return jsonify({
                'success': False,
                'message': f'文件解析失败: {str(e)}'
            })
        
        api_key = current_app.config.get('ZHIPU_API_KEY', '')
        if not api_key:
            return jsonify({
                'success': False,
                'message': '系统未配置API_KEY，请联系管理员'
            })
        
        if not GLM_SUPPORT:
            return jsonify({
                'success': False,
                'message': 'GLM SDK未安装，请联系管理员'
            })
        
        try:
            print("[简历解析] 开始调用GLM-4.7-FLASH...")
            glm_response = extract_resume_info_with_glm(text, api_key)
            print(f"[简历解析] GLM返回: {glm_response[:200]}..." if len(glm_response) > 200 else f"[简历解析] GLM返回: {glm_response}")
            
            extracted_data = parse_glm_response(glm_response)
            form_data = convert_to_form_data(extracted_data)
            
            return jsonify({
                'success': True,
                'message': '解析成功',
                'data': form_data
            })
            
        except Exception as e:
            error_msg = str(e)
            print(f"[简历解析] GLM调用异常: {error_msg}")
            
            if 'JSON' in error_msg or '解析' in error_msg:
                try:
                    print("[简历解析] 尝试重新调用GLM...")
                    glm_response = extract_resume_info_with_glm(text, api_key)
                    extracted_data = parse_glm_response(glm_response)
                    form_data = convert_to_form_data(extracted_data)
                    
                    return jsonify({
                        'success': True,
                        'message': '解析成功',
                        'data': form_data
                    })
                except Exception as retry_e:
                    print(f"[简历解析] 重试失败: {str(retry_e)}")
            
            return jsonify({
                'success': False,
                'message': f'AI解析失败: {error_msg}'
            })
            
    except Exception as e:
        print(f"[简历解析] 未捕获异常: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'服务器内部错误: {str(e)}'
        })
